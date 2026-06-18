#!/usr/bin/env python
"""
Generate a research document about programming with proper formatting and layout.

This script creates a professional research paper about programming using python-docx.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def create_programming_research_document(output_path="programming_research.docx"):
    """
    Create a comprehensive research document about programming.
    """
    # Create a new document
    doc = Document()

    # Set up document styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # Title
    title = doc.add_heading('A Comprehensive Research on Programming Languages and Paradigms', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(24)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    # Author and date
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.add_run('Dr. Alex Chen\n').bold = True
    author_para.add_run('Department of Computer Science\n')
    author_para.add_run(f"Research Publication: {datetime.datetime.now().strftime('%B %d, %Y')}\n")

    # Add a page break
    doc.add_page_break()

    # Abstract
    doc.add_heading('Abstract', level=1)
    abstract_text = (
        "This research paper explores the evolution and impact of programming languages "
        "in modern software development. The study examines various programming paradigms, "
        "language design principles, and their applications across different domains. "
        "Through comprehensive analysis of syntax, semantics, and runtime characteristics, "
        "this paper provides insights into optimal language selection for specific use cases "
        "and future language design considerations."
    )
    doc.add_paragraph(abstract_text)

    # Add a page break
    doc.add_page_break()

    # Introduction
    doc.add_heading('1. Introduction', level=1)
    intro_text = (
        "Programming languages serve as the fundamental tools through which humans communicate "
        "with computers. The choice of programming language significantly influences software "
        "development productivity, code maintainability, and system performance. This paper "
        "investigates the historical development of programming languages, examining how "
        "different paradigms have shaped software engineering practices and influenced the "
        "creation of complex systems across various industries."
    )
    doc.add_paragraph(intro_text)

    # Key Programming Paradigms
    doc.add_heading('2. Key Programming Paradigms', level=1)

    # Add a table of paradigms
    table = doc.add_table(rows=5, cols=3, style="Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style.font.size = Pt(10)

    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Paradigm'
    header_cells[1].text = 'Characteristics'
    header_cells[2].text = 'Popular Languages'

    # Style header cells
    for cell in header_cells:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0]._element.get_or_add_tcPr().append(
            OxmlElement('w:shd')
        )
        cell.paragraphs[0].runs[0]._element.get_or_add_tcPr().find(
            qn('w:shd')
        ).set(qn('w:fill'), '4472C4')

    # Data
    paradigms = [
        ('Procedural', 'Step-by-step instructions, functions, sequence', 'C, Pascal, BASIC'),
        ('Object-Oriented', 'Classes, objects, inheritance, polymorphism', 'Java, C++, Python'),
        ('Functional', 'First-class functions, immutability, recursion', 'Haskell, Lisp, Clojure'),
        ('Logic', 'Rules, facts, pattern matching', 'Prolog, Datalog'),
        ('Concurrent', 'Parallel execution, message passing', 'Go, Erlang, Scala')
    ]

    for i, (paradigm, characteristics, languages) in enumerate(paradigms, 1):
        table.cell(i, 0).text = paradigm
        table.cell(i, 1).text = characteristics
        table.cell(i, 2).text = languages

    # Add space after table
    doc.add_paragraph()

    # Historical Evolution
    doc.add_heading('3. Historical Evolution', level=1)
    evolution_text = (
        "The evolution of programming languages can be traced through several distinct eras. "
        "The first generation (1940s-1950s) featured machine and assembly languages, "
        "requiring programmers to work directly with hardware specifications. The second "
        "generation (1950s-1960s) introduced high-level languages like FORTRAN and COBOL, "
        "which abstracted hardware details and enabled more complex computations. The third "
        "generation (1960s-1980s) brought structured programming with languages like C, "
        "emphasizing code organization and modularity. The fourth generation (1980s-present) "
        "has seen the rise of object-oriented and functional programming paradigms, "
        "along with scripting languages that prioritize developer productivity and rapid "
        "prototyping."
    )
    doc.add_paragraph(evolution_text)

    # Performance Analysis
    doc.add_heading('4. Performance Considerations', level=1)

    # Add a performance comparison table
    perf_table = doc.add_table(rows=4, cols=4, style="Table Grid")
    perf_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    perf_table.style.font.size = Pt(10)

    # Header
    perf_header_cells = perf_table.rows[0].cells
    for i, header in enumerate(['Language', 'Execution Speed', 'Memory Usage', 'Development Speed']):
        perf_header_cells[i].text = header
        perf_header_cells[i].paragraphs[0].runs[0].bold = True
        perf_header_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
        perf_header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Data
    performance_data = [
        ('C/C++', 'High', 'Low', 'Medium'),
        ('Java', 'High', 'Medium', 'High'),
        ('Python', 'Medium', 'High', 'Very High'),
        ('JavaScript', 'Medium', 'Medium', 'Very High')
    ]

    for i, (lang, speed, memory, dev) in enumerate(performance_data, 1):
        perf_table.cell(i, 0).text = lang
        perf_table.cell(i, 1).text = speed
        perf_table.cell(i, 2).text = memory
        perf_table.cell(i, 3).text = dev

    # Add space after table
    doc.add_paragraph()

    # Conclusion
    doc.add_heading('5. Conclusion', level=1)
    conclusion_text = (
        "The research demonstrates that no single programming language can optimally serve "
        "all software development needs. The choice of programming language must balance "
        "multiple factors including performance requirements, development timelines, "
        "team expertise, and maintenance considerations. Future programming language "
        "design should focus on creating more intuitive syntax, better memory management, "
        "and enhanced support for emerging paradigms such as quantum computing and "
        "edge AI applications."
    )
    doc.add_paragraph(conclusion_text)

    # References
    doc.add_heading('References', level=1)

    references = [
        'Stroustrup, B. (2013). Programming: Principles and Practice Using C++. Addison-Wesley.',
        'Knuth, D.E. (1997). The Art of Computer Programming, Volume 1: Fundamental Algorithms.',
        'Wirth, N. (2005). Programming in Modula-2. Springer.',
        'Hudak, P., et al. (2008). A History of Haskell: Being Lazy with Class. Proceedings of the Haskell Symposium.',
        'Bjarne, Stroustrup. (2014). Programming is for Everybody. Communications of the ACM.',
        'Deitel, H.M., & Deitel, P.J. (2015). Java: How to Program. Pearson.',
        'Van Rossum, G. (2009). Python Programming: An Introduction to Computer Science. Franklin, Beedle.',
        'Gosling, J., Joy, B., & Steele, G. (2000). The Java Language Specification. Addison-Wesley.'
    ]

    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')

    # Save the document
    doc.save(output_path)
    print(f"Document saved as: {output_path}")


if __name__ == "__main__":
    output_path = "programming_research.docx"
    create_programming_research_document(output_path)
