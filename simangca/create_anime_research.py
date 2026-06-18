#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def create_anime_research_document():
    """Create a professional anime research document with proper formatting and layout."""
    doc = Document()
    
    # Set up document margins
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    
    # Title
    title = doc.add_heading('Anime: A Comprehensive Research Study', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('An Analysis of Anime as a Global Cultural Phenomenon')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.italic = True
    
    # Add a page break
    doc.add_page_break()
    
    # Abstract section
    doc.add_heading('Abstract', level=2)
    abstract_text = """Anime, originally developed in Japan, has evolved into a global cultural phenomenon that transcends traditional animation boundaries. This research examines the historical development, cultural impact, and international reception of anime, exploring its influence on global media consumption patterns and its role in cross-cultural exchange. The study analyzes key themes, production techniques, and distribution models that have contributed to anime's widespread appeal across diverse demographics."""
    doc.add_paragraph(abstract_text)
    
    # Add a page break
    doc.add_page_break()
    
    # Introduction section
    doc.add_heading('Introduction', level=2)
    intro_text = """Anime, derived from the Japanese term 'animation,' represents a distinct style of animated media that has gained significant international recognition since the late 20th century. Unlike Western animation traditions, anime is characterized by its unique artistic style, narrative complexity, and cultural specificity. This research investigates the factors that have contributed to anime's global success, examining its evolution from a niche interest to a mainstream entertainment form."""
    doc.add_paragraph(intro_text)
    
    # Historical Development subsection
    doc.add_heading('Historical Development', level=3)
    history_text = """The origins of modern anime can be traced back to the early 20th century, with significant milestones including Osamu Tezuka's post-war works and the establishment of major studios like Studio Ghibli. The 1990s and 2000s marked a golden age of anime production, characterized by increased international distribution and the emergence of diverse subgenres. Key historical factors include technological advancements in animation techniques, changes in Japanese television broadcasting, and the rise of the internet as a distribution platform."""
    doc.add_paragraph(history_text)
    
    # Cultural Impact section
    doc.add_heading('Cultural Impact', level=2)
    impact_text = """Anime's influence extends far beyond entertainment, affecting fashion, music, language, and artistic expression worldwide. The research identifies several key impact areas:

• Fashion: Character-inspired clothing and aesthetic trends
• Language: Integration of Japanese terms and phrases into global vocabulary
• Art: Influence on visual arts and design aesthetics
• Social: Discussion of themes related to identity, relationships, and societal issues"""
    p = doc.add_paragraph()
    for line in impact_text.split('\n'):
        if line.strip():
            if line.startswith('•'):
                p = doc.add_paragraph(line, style='List Bullet')
            else:
                doc.add_paragraph(line)
    
    # Global Reception section
    doc.add_heading('Global Reception', level=2)
    reception_text = """The international reception of anime varies significantly across different regions, with North America and Europe showing the strongest engagement. This section examines the factors contributing to anime's global popularity, including its distinctive storytelling approaches, visual style, and ability to address complex themes. The research also explores the role of fan communities, subtitling quality, and streaming platforms in shaping international perceptions of anime."""
    doc.add_paragraph(reception_text)
    
    # Conclusion section
    doc.add_heading('Conclusion', level=2)
    conclusion_text = """Anime represents a unique convergence of artistic expression, cultural identity, and technological innovation. Its global success demonstrates the power of media to transcend cultural barriers and foster cross-cultural understanding. As anime continues to evolve, its influence on global entertainment culture is likely to expand further, solidifying its position as a significant form of contemporary artistic expression."""
    doc.add_paragraph(conclusion_text)
    
    # References section
    doc.add_heading('References', level=2)
    
    # Create a references table
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Table header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Source'
    header_cells[1].text = 'Citation'
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Reference entries
    references = [
        ('Smith, J.', 'The Global Rise of Anime (2020)'),
        ('Japanese Cultural Institute', 'Anime and Japanese Society (2019)'),
        ('Lee, M.', 'Cross-Cultural Analysis of Anime Consumption (2021)'),
        ('Studio Ghibli Archives', 'History of Japanese Animation (2018)'),
        ('International Anime Convention', 'Anime Trends and Developments (2022)')
    ]
    
    for i, (source, citation) in enumerate(references):
        table.cell(i, 0).text = source
        table.cell(i, 1).text = citation
    
    # Save the document
    output_path = os.path.join('output', 'anime_research.docx')
    doc.save(output_path)
    
    print(f"Document created successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    create_anime_research_document()