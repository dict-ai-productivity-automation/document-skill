import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_document():
    doc = Document()
    
    # Header
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = "Research Report: Mythological Entities"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Title
    title = doc.add_heading("Mythical Creatures: A Cultural Study", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_heading("1. Introduction", level=1)
    p_intro = doc.add_paragraph("Mythical creatures have captivated human imagination for centuries. They appear in legends, folklore, and mythologies across diverse cultures worldwide. This document explores a few of the most prominent mythical beings, their origins, and their cultural significance.")
    p_intro.paragraph_format.first_line_indent = Inches(0.5)

    # Dragons
    doc.add_heading("2. Dragons", level=1)
    p_dragon = doc.add_paragraph("Dragons are perhaps the most universally recognized mythical creatures. ")
    run = p_dragon.add_run("In Western traditions")
    run.bold = True
    p_dragon.add_run(", they are often depicted as fearsome, fire-breathing, winged reptiles. ")
    run2 = p_dragon.add_run("In Eastern traditions")
    run2.bold = True
    p_dragon.add_run(", such as in Chinese mythology, they are revered as benevolent symbols of power, water, and good fortune.")
    
    # Characteristics List
    doc.add_paragraph("Key traits of Dragons:", style="List Bullet")
    doc.add_paragraph("Fire-breathing capabilities (Western)", style="List Bullet")
    doc.add_paragraph("Association with water and weather (Eastern)", style="List Bullet")
    doc.add_paragraph("Serpentine or reptilian bodies", style="List Bullet")

    # Unicorns
    doc.add_heading("3. Unicorns", level=1)
    p_unicorn = doc.add_paragraph("A symbol of purity and grace, the unicorn is typically described as a white horse-like creature with a single, spiraling horn projecting from its forehead. Folklore suggests that only the pure of heart could ever tame a unicorn.")
    p_unicorn.paragraph_format.first_line_indent = Inches(0.5)

    # Phoenix
    doc.add_heading("4. Phoenix", level=1)
    p_phoenix = doc.add_paragraph("The phoenix is an immortal bird associated with Greek mythology that cyclically regenerates or is otherwise born again. Associated with the sun, a phoenix obtains new life by arising from the ashes of its predecessor. It is a powerful symbol of rebirth and resilience.")
    p_phoenix.paragraph_format.first_line_indent = Inches(0.5)

    # Comparison Table
    doc.add_heading("5. Quick Comparison", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Creature'
    hdr_cells[1].text = 'Primary Domain'
    hdr_cells[2].text = 'Symbolism'

    data = [
        ('Dragon', 'Sky / Water', 'Power, Wisdom, Destruction'),
        ('Unicorn', 'Forests', 'Purity, Grace'),
        ('Phoenix', 'Fire / Sky', 'Rebirth, Immortality')
    ]

    for creature, domain, symbolism in data:
        row_cells = table.add_row().cells
        row_cells[0].text = creature
        row_cells[1].text = domain
        row_cells[2].text = symbolism

    # Conclusion
    doc.add_heading("6. Conclusion", level=1)
    p_conclusion = doc.add_paragraph("The enduring appeal of mythical creatures lies in what they represent. They embody human fears, hopes, values, and the sheer power of storytelling. While they may not exist in the physical world, they certainly live on in our cultural legacy.")
    p_conclusion.paragraph_format.first_line_indent = Inches(0.5)

    # Save document
    if not os.path.exists('output'):
        os.makedirs('output')
    
    doc.save('output/Mythical_Creatures_Research.docx')
    print("Document generated successfully at: output/Mythical_Creatures_Research.docx")

if __name__ == "__main__":
    create_document()
