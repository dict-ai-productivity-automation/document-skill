from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from datetime import datetime

# Create a new document
doc = Document()

# Set document properties
doc.core.title = "Greek Methodologies Research"
doc.core.author = "Research Team"
doc.core.subject = "Investigation of Ancient Greek Research Methods"
doc.core.keywords = "Greek, methodologies, research, philosophy, science"
doc.core.category = "Academic Research"
doc.core.comments = "Generated research document on Greek methodologies"
doc.core.created = datetime.now()

# Add title with custom formatting
title = doc.add_heading('Greek Methodologies in Ancient Research', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

# Add subtitle
subtitle = doc.add_paragraph('An Exploration of Systematic Inquiry and Knowledge Acquisition')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(12)
    run.font.italic = True

# Add page break
doc.add_page_break()

# Add introduction section
intro_heading = doc.add_heading('Introduction', level=1)
intro_heading.runs[0].font.size = Pt(18)
intro_heading.runs[0].font.bold = True

intro_text = doc.add_paragraph()
intro_text.add_run('The ancient Greeks established foundational methodologies that continue to influence modern scientific inquiry. ')
intro_text.add_run('Their systematic approach to knowledge acquisition encompassed philosophical, mathematical, and empirical investigation. ', bold=True)
intro_text.add_run('This research examines the key methodologies that defined Greek scholarly practices.')

# Add section break
doc.add_page_break()

# Add Greek Philosophical Methodologies section
phil_heading = doc.add_heading('Greek Philosophical Methodologies', level=1)
phil_heading.runs[0].font.size = Pt(18)
phil_heading.runs[0].font.bold = True

# Create a table of philosophical methodologies
method_table = doc.add_table(rows=4, cols=3)
method_table.style = 'Table Grid'
method_table.cell(0, 0).text = 'Methodology'
method_table.cell(0, 1).text = 'Key Figures'
method_table.cell(0, 2).text = 'Core Principles'

method_table.cell(1, 0).text = 'Dialectic Method'
method_table.cell(1, 1).text = 'Socrates, Plato'
method_table.cell(1, 2).text = 'Thesis-Antithesis-Synthesis'

method_table.cell(2, 0).text = 'Empirical Observation'
method_table.cell(2, 1).text = 'Aristotle'
method_table.cell(2, 2).text = 'Systematic Data Collection'

method_table.cell(3, 0).text = 'Mathematical Reasoning'
method_table.cell(3, 1).text = 'Pythagoras, Euclid'
method_table.cell(3, 2).text = 'Axiomatic Proofs'

# Add section break
doc.add_page_break()

# Add Mathematical Methodologies section
math_heading = doc.add_heading('Mathematical Methodologies', level=1)
math_heading.runs[0].font.size = Pt(18)
math_heading.runs[0].font.bold = True

math_text = doc.add_paragraph()
math_text.add_run('Greek mathematics was characterized by rigorous logical deduction and formal proof. ')
math_text.add_run('The axiomatic method, developed through Euclid\'s Elements, established a template for mathematical reasoning that persists today. ')
math_text.add_run('Key mathematical methodologies included the use of definitions, postulates, and common notions as foundations for proofs.', bold=True)

# Add a bullet list of mathematical contributions
math_heading2 = doc.add_heading('Major Mathematical Contributions', level=2)
bullet_style = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)

bullet1 = doc.add_paragraph('Euclidean Geometry - The systematic study of plane and solid geometry', style='List Bullet')
bullet1.runs[0].font.size = Pt(11)

bullet2 = doc.add_paragraph('Number Theory - Study of integer properties and relationships', style='List Bullet')
bullet2.runs[0].font.size = Pt(11)

bullet3 = doc.add_paragraph('Geometric Algebra - Integration of algebraic concepts with geometric representation', style='List Bullet')
bullet3.runs[0].font.size = Pt(11)

# Add section break
doc.add_page_break()

# Add Scientific Methodologies section
science_heading = doc.add_heading('Scientific Methodologies', level=1)
science_heading.runs[0].font.size = Pt(18)
science_heading.runs[0].font.bold = True

science_text = doc.add_paragraph()
science_text.add_run('Aristotelian science represents one of the earliest systematic attempts at empirical investigation. ')
science_text.add_run('Through his work in the Organon, Aristotle established methodologies for logical analysis and categorization of knowledge. ')
science_text.add_run('His approach combined observation with logical deduction, creating a framework that influenced scientific inquiry for centuries.', bold=True)

# Add a table of scientific methods
science_table = doc.add_table(rows=3, cols=2)
science_table.style = 'Table Grid'
science_table.cell(0, 0).text = 'Method'
science_table.cell(0, 1).text = 'Application'

science_table.cell(1, 0).text = 'Systematic Classification'
science_table.cell(1, 1).text = 'Biology, Zoology'

science_table.cell(2, 0).text = 'Causal Explanation'
science_table.cell(2, 1).text = 'Physics, Metaphysics'

# Add section break
doc.add_page_break()

# Add Comparative Analysis section
comp_heading = doc.add_heading('Comparative Analysis: Greek vs. Modern Methodologies', level=1)
comp_heading.runs[0].font.size = Pt(18)
comp_heading.runs[0].font.bold = True

# Create comparison table
comp_table = doc.add_table(rows=4, cols=2)
comp_table.style = 'Table Grid'
comp_table.cell(0, 0).text = 'Aspect'
comp_table.cell(0, 1).text = 'Greek Methodology'

comp_table.cell(1, 0).text = 'Foundation'
comp_table.cell(1, 1).text = 'Philosophical Inquiry'

comp_table.cell(2, 0).text = 'Evidence'
comp_table.cell(2, 1).text = 'Observation + Reason'

comp_table.cell(3, 0).text = 'Validation'
comp_table.cell(3, 1).text = 'Logical Consistency'

# Add section break
doc.add_page_break()

# Add Conclusion section
concl_heading = doc.add_heading('Conclusion', level=1)
concl_heading.runs[0].font.size = Pt(18)
concl_heading.runs[0].font.bold = True

concl_text = doc.add_paragraph()
concl_text.add_run('The Greek methodologies established enduring principles that continue to shape modern research practices. ')
concl_text.add_run('From the dialectic method to empirical observation, Greek scholars developed systematic approaches to knowledge acquisition that emphasized logical reasoning, evidence-based inquiry, and structured argumentation. ')
concl_text.add_run('These methodologies formed the foundation upon which modern scientific and academic practices are built.', bold=True)

# Add a final summary box
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls

# Add a paragraph with a distinctive style
summary = doc.add_paragraph()
summary.style = 'Intense Quote'
summary.add_run('Key Takeaways:\n')
summary.add_run('• Systematic methodology over anecdotal evidence\n')
summary.add_run('• Logical deduction as primary validation tool\n')
summary.add_run('• Integration of observation and reason\n')
summary.add_run('• Emphasis on clear definitions and terminology\n')
summary.add_run('• Structured argumentation and proof\n')

# Add footer with page numbers
section = doc.sections[-1]
footer = section.footer
footer_para = footer.paragraphs[0]
footer_para.text = f"Page {len(doc.paragraphs)}"
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in footer_para.runs:
    run.font.size = Pt(10)

# Save the document
doc.save('greek_methodologies_research.docx')
print("Document 'greek_methodologies_research.docx' created successfully!")
print(f"Document contains {len(doc.paragraphs)} paragraphs")
print(f"Document contains {len(doc.tables)} tables")