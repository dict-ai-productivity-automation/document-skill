"""DOCX document builder for academic papers.

This module builds DOCX documents from parsed Markdown content with
proper academic formatting, styles, and layout.
"""

import os
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def build_docx(
    data: Dict[str, Any],
    output_path: str,
    font: str = "Times New Roman",
    font_size: int = 12,
    heading_font: str = "Times New Roman",
    line_spacing: float = 1.5,
    margin: float = 1.0,
    page_size: str = "letter",
    include_toc: bool = True,
    include_page_numbers: bool = True,
    include_running_head: bool = False,
    **style_config
) -> None:
    """Build DOCX document from parsed data.

    Args:
        data: Parsed document data
        output_path: Output file path
        font: Body font name
        font_size: Body font size (pt)
        heading_font: Heading font name
        line_spacing: Line spacing
        margin: Page margins (inches)
        page_size: Page size ("letter" or "a4")
        include_toc: Include table of contents
        include_page_numbers: Include page numbers
        include_running_head: Include running head
        **style_config: Additional style configuration
    """
    # Create new document
    doc = Document()

    # Configure page setup
    _configure_page_setup(doc, margin, page_size)

    # Add title page
    _add_title_page(doc, data, font, font_size)

    # Add table of contents
    if include_toc:
        _add_table_of_contents(doc, data, font, font_size)

    # Add main content
    _add_main_content(doc, data, font, font_size, heading_font, line_spacing)

    # Add references section
    _add_references_section(doc, data, font, font_size)

    # Add page numbers and running head
    if include_page_numbers:
        _add_page_numbers(doc)
    if include_running_head:
        _add_running_head(doc, data)

    # Save document
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def _configure_page_setup(doc: Document, margin: float, page_size: str) -> None:
    """Configure page setup (margins, size, etc.).

    Args:
        doc: Document object
        margin: Page margins (inches)
        page_size: Page size ("letter" or "a4")
    """
    section = doc.sections[0]

    # Set margins
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

    # Set page size
    if page_size.lower() == "a4":
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    else:  # letter
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)


def _add_title_page(doc: Document, data: Dict[str, Any], font: str, font_size: int) -> None:
    """Add title page to document.

    Args:
        doc: Document object
        data: Document data
        font: Font name
        font_size: Font size (pt)
    """
    # Add title
    title = doc.add_paragraph(data["title"], style="Title")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].font.name = font
    title.runs[0].font.size = Pt(font_size + 8)

    # Add author
    author = doc.add_paragraph(f"by {data['author']}", style="Subtitle")
    author.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    author.runs[0].font.name = font
    author.runs[0].font.size = Pt(font_size)

    # Add date
    date = doc.add_paragraph(f"{data['date']}", style="Subtitle")
    date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date.runs[0].font.name = font
    date.runs[0].font.size = Pt(font_size)

    # Add abstract box
    if data.get("abstract"):
        abstract_title = doc.add_paragraph("Abstract", style="Heading 1")
        abstract_title.runs[0].font.name = font
        abstract_title.runs[0].font.size = Pt(font_size + 2)

        abstract_text = doc.add_paragraph(data["abstract"])
        abstract_text.runs[0].font.name = font
        abstract_text.runs[0].font.size = Pt(font_size)

        # Add some spacing after abstract
        doc.add_paragraph()


def _add_table_of_contents(doc: Document, data: Dict[str, Any], font: str, font_size: int) -> None:
    """Add table of contents to document.

    Args:
        doc: Document object
        data: Document data
        font: Font name
        font_size: Font size (pt)
    """
    # Add TOC heading
    toc_heading = doc.add_paragraph("Table of Contents", style="Heading 1")
    toc_heading.runs[0].font.name = font
    toc_heading.runs[0].font.size = Pt(font_size + 2)

    # Add TOC entries
    for section in data["sections"]:
        if section["level"] == 1:
            # Level 1 heading (main section)
            toc_entry = doc.add_paragraph(f"{section['title']}", style="List Paragraph")
            toc_entry.runs[0].font.name = font
            toc_entry.runs[0].font.size = Pt(font_size)

            # Add sub-sections
            for sub_section in data["sections"]:
                if sub_section["level"] == 2 and sub_section["title"].startswith(section["title"]):
                    sub_entry = doc.add_paragraph(
                        f"  {sub_section['title']}",
                        style="List Paragraph"
                    )
                    sub_entry.runs[0].font.name = font
                    sub_entry.runs[0].font.size = Pt(font_size - 2)

    # Add spacing after TOC
    doc.add_paragraph()


def _add_main_content(
    doc: Document,
    data: Dict[str, Any],
    font: str,
    font_size: int,
    heading_font: str,
    line_spacing: float
) -> None:
    """Add main content to document.

    Args:
        doc: Document object
        data: Document data
        font: Font name
        font_size: Font size (pt)
        heading_font: Heading font name
        line_spacing: Line spacing
    """
    for section in data["sections"]:
        # Add section heading
        if section["level"] == 1:
            heading = doc.add_paragraph(section["title"], style="Heading 1")
            heading.runs[0].font.name = heading_font
            heading.runs[0].font.size = Pt(font_size + 4)
            heading.runs[0].bold = True

        elif section["level"] == 2:
            heading = doc.add_paragraph(section["title"], style="Heading 2")
            heading.runs[0].font.name = heading_font
            heading.runs[0].font.size = Pt(font_size + 2)
            heading.runs[0].bold = True

        elif section["level"] == 3:
            heading = doc.add_paragraph(section["title"], style="Heading 3")
            heading.runs[0].font.name = heading_font
            heading.runs[0].font.size = Pt(font_size + 1)

        # Add section content
        if section["content"]:
            # Split content into paragraphs
            paragraphs = section["content"].split('\n\n')
            for paragraph_text in paragraphs:
                if paragraph_text.strip():
                    paragraph = doc.add_paragraph(paragraph_text)
                    paragraph.runs[0].font.name = font
                    paragraph.runs[0].font.size = Pt(font_size)

        # Add spacing after section
        doc.add_paragraph()


def _add_references_section(doc: Document, data: Dict[str, Any], font: str, font_size: int) -> None:
    """Add references section to document.

    Args:
        doc: Document object
        data: Document data
        font: Font name
        font_size: Font size (pt)
    """
    # Add references heading
    ref_heading = doc.add_paragraph("References", style="Heading 1")
    ref_heading.runs[0].font.name = font
    ref_heading.runs[0].font.size = Pt(font_size + 4)
    ref_heading.runs[0].bold = True

    # Add references
    for i, ref in enumerate(data["references"], 1):
        ref_text = f"{i}. {ref['text']}"
        ref_paragraph = doc.add_paragraph(ref_text)
        ref_paragraph.runs[0].font.name = font
        ref_paragraph.runs[0].font.size = Pt(font_size)

    # Add spacing after references
    doc.add_paragraph()


def _add_page_numbers(doc: Document) -> None:
    """Add page numbers to document.

    Args:
        doc: Document object
    """
    section = doc.sections[0]

    # Add footer with page numbers
    footer = section.footer
    footer.paragraphs[0].text = "Page " + str(doc.add_page_number())
    footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def _add_running_head(doc: Document, data: Dict[str, Any]) -> None:
    """Add running head to document.

    Args:
        doc: Document object
        data: Document data
    """
    section = doc.sections[0]

    # Add header with running head
    header = section.header
    running_head_text = f"{data['title']} - {data['author']}"
    header.paragraphs[0].text = running_head_text
    header.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def add_table(doc: Document, data: List[List[str]], title: Optional[str] = None) -> None:
    """Add a table to the document.

    Args:
        doc: Document object
        data: Table data as list of rows (each row is a list of cells)
        title: Optional table title
    """
    if title:
        title_paragraph = doc.add_paragraph(title)
        title_paragraph.runs[0].bold = True

    table = doc.add_table(rows=len(data), cols=len(data[0]) if data else 0)
    table.style = "Table Grid"

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = str(cell_data)

    # Add caption if title provided
    if title:
        caption = doc.add_paragraph(f"Figure: {title}")
        caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def add_figure(doc: Document, image_path: str, caption: str = "") -> None:
    """Add a figure to the document.

    Args:
        doc: Document object
        image_path: Path to image file
        caption: Figure caption
    """
    if not os.path.exists(image_path):
        print(f"Warning: Image file not found: {image_path}")
        return

    # Add figure
    doc.add_picture(image_path, width=Inches(6.0))

    # Add caption
    if caption:
        caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        caption_paragraph.runs[0].italic = True


def add_code_block(doc: Document, code: str, language: str = "") -> None:
    """Add a code block to the document.

    Args:
        doc: Document object
        code: Code content
        language: Programming language
    """
    # Add code block paragraph
    code_paragraph = doc.add_paragraph()

    # Create run for code
    run = code_paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(10)

    # Add shaded background (simulated with formatting)
    shading = run._element.rPr
    shading.set(qn('w:shd'), 'clear')

    # Add language label if provided
    if language:
        lang_paragraph = doc.add_paragraph(f"Language: {language}")
        lang_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        lang_paragraph.runs[0].font.size = Pt(8)
        lang_paragraph.runs[0].italic = True


def add_list(doc: Document, items: List[str], ordered: bool = False) -> None:
    """Add a list to the document.

    Args:
        doc: Document object
        items: List of list items
        ordered: Whether to create an ordered list
    """
    if ordered:
        for i, item in enumerate(items, 1):
            doc.add_paragraph(f"{i}. {item}")
    else:
        for item in items:
            doc.add_paragraph(f"• {item}")
